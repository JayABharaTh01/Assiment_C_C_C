"""
data_read.py

Reads all supported documents from the data directory and converts them
into Document objects.

Supported Formats:
    - PDF
    - CSV
    - JSON
    - DOCX

Author: Jaya Bharath
"""

from pathlib import Path
import json

import fitz  # PyMuPDF
import pandas as pd
from docx import Document as DocxDocument

from models.document import Document
from utils.config import DATA_FOLDER, SUPPORTED_EXTENSIONS


class DataReader:
    """
    Reads all supported files from the configured data folder.
    """

    def __init__(self):
        self.data_folder = Path(DATA_FOLDER)

        if not self.data_folder.exists():
            raise FileNotFoundError(
                f"Data folder not found: {self.data_folder}"
            )

    # ==========================================================
    # PDF Reader
    # ==========================================================

    def _read_pdf(self, file_path: Path) -> str:
        """
        Read text from a PDF file.
        """

        text = []

        with fitz.open(file_path) as pdf:

            for page in pdf:
                text.append(page.get_text())

        return "\n".join(text)

    # ==========================================================
    # CSV Reader
    # ==========================================================

    def _read_csv(self, file_path: Path) -> str:
        """
        Read CSV and convert it into plain text.
        """

        df = pd.read_csv(file_path)

        return df.to_string(index=False)

    # ==========================================================
    # JSON Reader
    # ==========================================================

    def _read_json(self, file_path: Path) -> str:
        """
        Read JSON and convert into formatted string.
        """

        with open(file_path, "r", encoding="utf-8") as file:
            data = json.load(file)

        return json.dumps(data, indent=2)

    # ==========================================================
    # DOCX Reader
    # ==========================================================

    def _read_docx(self, file_path: Path) -> str:
        """
        Read Microsoft Word document.
        """

        document = DocxDocument(file_path)

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        return "\n".join(paragraphs)

    # ==========================================================
    # Detect Reader
    # ==========================================================

    def _read_file(self, file_path: Path) -> str:
        """
        Automatically detect file type and call the correct reader.
        """

        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return self._read_pdf(file_path)

        elif extension == ".csv":
            return self._read_csv(file_path)

        elif extension == ".json":
            return self._read_json(file_path)

        elif extension == ".docx":
            return self._read_docx(file_path)

        raise ValueError(
            f"Unsupported file type: {extension}"
        )

    # ==========================================================
    # Read All Files
    # ==========================================================

    def read_documents(self):
        """
        Read every supported document inside the data folder.

        Returns
        -------
        list[Document]
        """

        documents = []

        for file_path in sorted(self.data_folder.iterdir()):

            if not file_path.is_file():
                continue

            if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue

            content = self._read_file(file_path)

            document = Document(
                filename=file_path.name,
                filetype=file_path.suffix.lower(),
                content=content
            )

            documents.append(document)

        return documents